import os
from docx import Document
import pandas as pd
from django.core.files.storage import FileSystemStorage

def convert_doc_to_excel(doc_file):
    fs = FileSystemStorage()
    filename = fs.save(doc_file.name, doc_file)
    file_path = fs.path(filename)  

    doc = Document(file_path)
    
    table_data = []
    for table in doc.tables:
        for row in table.rows:
            row_data = []
            for cell in row.cells:
                row_data.append(cell.text.strip())  
            table_data.append(row_data)  
    if not table_data:
        table_data = [[para.text] for para in doc.paragraphs]
    if table_data:
        df = pd.DataFrame(table_data[1:], columns=table_data[0])
    else:
        df = pd.DataFrame(table_data, columns=["Text"])

    # Save the DataFrame as an Excel file
    excel_filename = os.path.splitext(filename)[0] + ".xlsx"
    excel_path = os.path.join('media', excel_filename)
    df.to_excel(excel_path, index=False)

    return excel_path
