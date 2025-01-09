import fitz  # PyMuPDF
from django.shortcuts import render
from .forms import UploadFileForm
from django.http import HttpResponse
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

def convert_pdf_to_word(request):
    if request.method == 'POST':
        form = UploadFileForm(request.POST, request.FILES)
        if form.is_valid():
            uploaded_file = request.FILES['file']
            doc = Document()
            pdf_document = fitz.open(stream=uploaded_file.read(), filetype="pdf")

            for page_num in range(len(pdf_document)):
                page = pdf_document.load_page(page_num)
                blocks = page.get_text("dict")["blocks"]

                for block in blocks:
                    para = doc.add_paragraph()
                    para_format = para.paragraph_format
                    para_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
                    para_format.left_indent = Pt(10)
                    para_format.space_before = Pt(5)
                    para_format.space_after = Pt(5)
                    para_format.line_spacing = 1.5

                    for line in block.get("lines", []):
                        for span in line.get("spans", []):
                            run = para.add_run(span["text"])
                            run.font.size = Pt(span["size"])
                            run.font.name = span["font"]

                            # Improved Color Handling
                            if "color" in span:
                                color_int = span['color']
                                r = (color_int >> 16) & 0xFF
                                g = (color_int >> 8) & 0xFF
                                b = color_int & 0xFF
                                run.font.color.rgb = RGBColor(r, g, b)

                            # Bold and Italics Detection
                            if "bold" in span["font"].lower():
                                run.bold = True
                            if "italic" in span["font"].lower():
                                run.italic = True

            response = HttpResponse(
                content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
            )
            response['Content-Disposition'] = 'attachment; filename="converted.docx"'
            doc.save(response)
            return response
    else:
        form = UploadFileForm()
    return render(request, 'uploadPDF.html', {'form': form})
